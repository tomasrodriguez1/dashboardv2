from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def create_proposal_docx():
    doc = Document()

    # --- Configuración de Estilos ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Estilo de Título 1
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(16)
    h1_font.color.rgb = RGBColor(0, 51, 102) # Azul Corporativo Oscuro
    h1_font.bold = True

    # Estilo de Título 2
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(14)
    h2_font.color.rgb = RGBColor(47, 84, 150) # Azul Medio
    h2_font.bold = True

    # Estilo de Título 3
    h3_style = doc.styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Arial'
    h3_font.size = Pt(12)
    h3_font.color.rgb = RGBColor(68, 84, 106) # Gris Azulado
    h3_font.bold = True

    # --- Funciones de Ayuda ---
    def add_page_number(run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def set_footer(doc):
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("Confidencial | Página ")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        add_page_number(paragraph.add_run())

    # --- PORTADA ---
    doc.add_paragraph("\n\n\n\n")
    title = doc.add_paragraph("Propuesta Técnica\nSkaba BI Solution")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = doc.styles['Title']
    
    subtitle = doc.add_paragraph("\nCliente: Skava\nProyecto: Implementación de Plataforma de Inteligencia de Negocios")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = doc.styles['Subtitle']

    doc.add_paragraph("\n\n\n\n\n\n\n\n")
    date_para = doc.add_paragraph(f"Fecha: {datetime.date.today().strftime('%d/%m/%Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- CONTROL DE VERSIONES ---
    doc.add_heading('Control de Versiones', level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Versión'
    hdr_cells[1].text = 'Fecha'
    hdr_cells[2].text = 'Autor'
    hdr_cells[3].text = 'Descripción'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = '1.0'
    row_cells[1].text = datetime.date.today().strftime('%d/%m/%Y')
    row_cells[2].text = 'Tomás Rodríguez'
    row_cells[3].text = 'Versión final para revisión del cliente.'

    doc.add_page_break()

    # --- CONTENIDO ---

    # 1. Resumen Ejecutivo
    doc.add_heading('1. Resumen ejecutivo', level=1)
    p = doc.add_paragraph("El objetivo principal de la presente propuesta es el diseño, construcción e implementación de 'Skaba BI Solution', una plataforma de inteligencia de negocios ejecutiva que integra la proyección comercial (ventas, oportunidades y fuerza de ventas) con la realidad operacional de mantenciones (órdenes de trabajo, consumo de repuestos y tiempos de entrega).")
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph("La solución permitirá a Skava anticipar la demanda de repuestos, detectar brechas entre la venta real y el ingreso potencial, y coordinar decisiones estratégicas entre las áreas comerciales y operacionales. Adicionalmente, se incorporará un Agente de Inteligencia Artificial (Skava AI) basado en tecnologías de Procesamiento de Lenguaje Natural (NLP) para facilitar la exploración de datos.")
    
    doc.add_paragraph("El proyecto se plantea como una plataforma web escalable, diseñada para soportar inicialmente hasta 20 usuarios ejecutivos/analistas, con capacidad de crecimiento proyectada. La ejecución se llevará a cabo en dos etapas: una fase inicial de MVP (6 semanas) y una fase de Despliegue y Cierre (4 semanas).")

    # 2. Contexto y objetivo del cliente
    doc.add_heading('2. Contexto y objetivo del cliente', level=1)
    doc.add_paragraph("Skava requiere disponer de una plataforma ejecutiva unificada que conecte la proyección comercial con la gestión de mantenimiento operacional. La necesidad central es fundamentar la toma de decisiones en datos confiables, poniendo foco en la trazabilidad de la información, la explicabilidad de los indicadores y la facilidad de uso para perfiles ejecutivos y analíticos.")

    # 3. Alcance
    doc.add_heading('3. Alcance', level=1)
    
    doc.add_heading('3.1 Alcance de la Solución', level=2)
    doc.add_paragraph("El servicio considera el diseño, desarrollo e implementación de un MVP de la plataforma Skaba BI Solution, abarcando los siguientes componentes:")
    
    doc.add_paragraph("1. Dashboards Ejecutivos (MVP):", style='List Bullet')
    doc.add_paragraph("Implementación de tableros web con visualizaciones, filtros y métricas para Ventas (desempeño y brechas), Fuerza de Ventas (ranking y cobertura) y Repuestos/Mantenciones (analítica de demanda).")
    
    doc.add_paragraph("2. Skava AI (MVP):", style='List Bullet')
    doc.add_paragraph("Asistente en lenguaje natural con capacidad RAG (Retrieval-Augmented Generation) sobre datos iniciales, explicabilidad de KPIs e insights automáticos básicos.")
    
    doc.add_paragraph("3. Datos y Gobierno:", style='List Bullet')
    doc.add_paragraph("Establecimiento de un modelo semántico inicial, catálogo de KPIs trazables, linaje de datos mínimo y reglas básicas de calidad.")
    
    doc.add_paragraph("4. Ingesta de Datos:", style='List Bullet')
    doc.add_paragraph("Estrategia dual: Carga controlada de datos (T+1) para la Etapa 1 y habilitación de conectores ETL/ELT para fuentes reales en la Etapa 2.")
    
    doc.add_paragraph("5. Seguridad y Control de Acceso:", style='List Bullet')
    doc.add_paragraph("Implementación de autenticación y autorización basada en roles (RBAC) y segmentación geográfica.")
    
    doc.add_paragraph("6. Observabilidad y Documentación:", style='List Bullet')
    doc.add_paragraph("Capa de monitoreo de salud de servicios, logs estructurados, documentación técnica final, runbooks y capacitación.")

    doc.add_paragraph("7. Mantención Mensual (Uso Razonable):", style='List Bullet')
    doc.add_paragraph("Incluye costos operativos de hosting, proveedor de LLM (hasta 200 interacciones/mes) y soporte operacional base.")

    doc.add_heading('3.2 Alcance por Etapa', level=2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Etapa"
    table.rows[0].cells[1].text = "Componentes Clave"
    
    table.rows[1].cells[0].text = "Etapa 1 – MVP"
    table.rows[1].cells[1].text = "- Dashboards de Ventas, Vendedores y Repuestos.\n- Skava AI sobre datos controlados (mock).\n- Modelo semántico y gobierno de datos base.\n- Seguridad RBAC."
    
    table.rows[2].cells[0].text = "Etapa 2 – Integraciones y Cierre"
    table.rows[2].cells[1].text = "- Integración de ingesta de datos con sistemas productivos.\n- Ajustes con datos reales y validación de linaje.\n- Documentación final y handover.\n- Soporte funcional."

    # 4. Solución propuesta
    doc.add_heading('4. Solución propuesta', level=1)
    
    doc.add_heading('4.1 Descripción de la solución', level=2)
    doc.add_paragraph("Skaba BI Solution se define como una plataforma web integrada que unifica la visión comercial y operacional. La solución prioriza la confiabilidad del dato mediante procesos de gobierno y trazabilidad, y facilita la interacción del usuario mediante interfaces visuales intuitivas y asistencia por Inteligencia Artificial.")

    doc.add_heading('4.2 Stack y contexto', level=2)
    doc.add_paragraph("La arquitectura tecnológica seleccionada asegura escalabilidad y mantenibilidad:")
    p = doc.add_paragraph()
    p.add_run("• Frontend: ").bold = True
    p.add_run("TypeScript, React, Vite.\n")
    p.add_run("• Backend: ").bold = True
    p.add_run("Node.js/TS.\n")
    p.add_run("• Datos: ").bold = True
    p.add_run("PostgreSQL, Vector DB.\n")
    p.add_run("• Orquestación e IA: ").bold = True
    p.add_run("n8n, Anthropic (LLM provider).\n")
    p.add_run("• Infraestructura: ").bold = True
    p.add_run("Render.")

    doc.add_heading('4.3 Arquitectura', level=2)
    
    doc.add_heading('4.3.1 Resumen Ejecutivo de Arquitectura', level=3)
    doc.add_paragraph("La arquitectura propuesta sigue un enfoque modular basado en servicios. El Frontend (React) consume APIs (REST/GraphQL) servidas por un Backend (Node.js). Los datos residen en un esquema analítico en PostgreSQL (capas raw/curated/semantic). La orquestación de datos se gestiona mediante n8n, permitiendo flexibilidad para cargas manuales en la Etapa 1 e integraciones automatizadas en la Etapa 2. El componente de IA utiliza una base de datos vectorial gestionada para implementar el patrón RAG.")

    doc.add_heading('4.3.2 Vista de Contexto (C4 Nivel 1)', level=3)
    doc.add_paragraph("El sistema interactúa con los siguientes actores y sistemas externos:")
    doc.add_paragraph("• Usuarios: Ejecutivos y Analistas.", style='List Bullet')
    doc.add_paragraph("• Sistemas Fuente: Archivos planos (E1) y ERP/CRM [PENDIENTE DE DEFINIR] (E2).", style='List Bullet')
    doc.add_paragraph("• Servicios Externos: Proveedor de LLM, Vector DB, Servicios de Correo.", style='List Bullet')

    doc.add_heading('4.3.3 Vista de Contenedores (C4 Nivel 2)', level=3)
    doc.add_paragraph("La solución se compone de los siguientes contenedores lógicos:")
    doc.add_paragraph("1. Frontend Web App (HTTPS/TLS).", style='List Bullet')
    doc.add_paragraph("2. Backend API (Auth, Business Logic).", style='List Bullet')
    doc.add_paragraph("3. Base de Datos Analítica (PostgreSQL).", style='List Bullet')
    doc.add_paragraph("4. Motor de Orquestación (n8n).", style='List Bullet')
    doc.add_paragraph("5. Base de Datos Vectorial.", style='List Bullet')

    doc.add_heading('4.3.4 Contenedores y responsabilidades', level=3)
    doc.add_paragraph("Se definen las responsabilidades principales:")
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Frontend Web"
    table.rows[0].cells[1].text = "Renderizado de dashboards, interfaz de chat, gestión de usuarios."
    table.rows[1].cells[0].text = "Backend API"
    table.rows[1].cells[1].text = "Lógica de negocio, autenticación, orquestación RAG."
    table.rows[2].cells[0].text = "PostgreSQL"
    table.rows[2].cells[1].text = "Almacenamiento analítico, vistas materializadas."
    table.rows[3].cells[0].text = "Metadata Service"
    table.rows[3].cells[1].text = "Catálogo de KPIs, linaje y calidad de datos."
    table.rows[4].cells[0].text = "n8n Orchestrator"
    table.rows[4].cells[1].text = "ETL/ELT, flujos de integración y workflows de calidad."
    table.rows[5].cells[0].text = "RAG Service"
    table.rows[5].cells[1].text = "Control de flujo de preguntas y respuestas con LLM."

    # 5. Requerimientos
    doc.add_heading('5. Requerimientos', level=1)
    
    doc.add_heading('5.1 Requisitos funcionales — Etapa 1 (MVP)', level=2)
    doc.add_paragraph("A. Datos y Gobierno:", style='List Bullet')
    doc.add_paragraph("Disponibilidad de modelo semántico, catálogo de KPIs, ingesta periódica controlada, calidad básica, seguridad RBAC y linaje mínimo.")
    doc.add_paragraph("B. Inteligencia Comercial (Dashboard Ventas):", style='List Bullet')
    doc.add_paragraph("Visión ejecutiva de ingresos, análisis de brecha comercial (Real vs Potencial), comparativos por país/mercado, análisis por categoría y filtros de segmentación.")
    doc.add_paragraph("C. Fuerza de Ventas:", style='List Bullet')
    doc.add_paragraph("Seguimiento de desempeño, cobertura, actividad comercial y normalización de comparativas.")
    doc.add_paragraph("D. Repuestos y Mantenciones:", style='List Bullet')
    doc.add_paragraph("Proyección básica de demanda, análisis jerárquico, señales de lead time y valorización estimada.")
    doc.add_paragraph("E. Skava AI:", style='List Bullet')
    doc.add_paragraph("Consulta en lenguaje natural, explicación de KPIs, insights automáticos básicos y escalamiento a soporte.")

    doc.add_heading('5.2 Requisitos funcionales — Etapa 2', level=2)
    doc.add_paragraph("F. Soporte Funcional y Analítico:", style='List Bullet')
    doc.add_paragraph("Gestión formal de incidencias, soporte nivelado y acompañamiento.")
    doc.add_paragraph("G. Integraciones:", style='List Bullet')
    doc.add_paragraph("Integración vía ETL/ELT con sistemas [PENDIENTE DE DEFINIR], validación/reconciliación de datos y robustez operativa.")
    doc.add_paragraph("H. Ajustes post-integración:", style='List Bullet')
    doc.add_paragraph("Calibración del modelo y tableros con datos productivos, ejecución de QA/UAT.")
    doc.add_paragraph("I. Cierre y Traspaso:", style='List Bullet')
    doc.add_paragraph("Entrega de documentación operativa, capacitación y plan de evolución.")

    doc.add_heading('5.3 Requisitos no funcionales (RNF)', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ["ID", "Categoría", "Objetivo"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    
    rnfs = [
        ("RNF-PERF-01", "Latencia", "Carga inicial ≤5s (p95); filtros ≤2s (p95)."),
        ("RNF-AV-01", "Disponibilidad", "Prod 99% horario comercial; MVP ≥95%."),
        ("RNF-DATA-01", "Frescura", "T+1 (MVP)."),
        ("RNF-SEC-01", "Seguridad", "Auth obligatoria, RBAC rol/país, HTTPS."),
        ("RNF-SCAL-01", "Escalabilidad", "Arquitectura preparada para ~100 usuarios.")
    ]
    for i, (rid, cat, obj) in enumerate(rnfs, 1):
        table.rows[i].cells[0].text = rid
        table.rows[i].cells[1].text = cat
        table.rows[i].cells[2].text = obj

    # 6. Plan de implementación
    doc.add_heading('6. Plan de implementación', level=1)
    
    doc.add_heading('6.1 Plan de sprints (alineado a inversión y SP)', level=2)
    doc.add_paragraph("El plan de trabajo se estructura en sprints quincenales, con una estimación referencial de capacidad.")
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ["Sprint", "Etapa", "Objetivo Principal"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        
    sprints = [
        ("S1 (W1-W2)", "Etapa 1", "Base de plataforma, Modelo semántico, Auth/RBAC."),
        ("S2 (W3-W4)", "Etapa 1", "Inteligencia Comercial y Tableros base."),
        ("S3 (W5-W6)", "Etapa 1", "Skava AI (MVP), Gobierno mínimo. Hito MVP."),
        ("S4 (W7-W8)", "Etapa 2", "Integraciones productivas (ETL/ELT), Validación."),
        ("S5 (W9-W10)", "Etapa 2", "Ajustes, UAT, Documentación y Cierre.")
    ]
    for i, row_data in enumerate(sprints, 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val

    doc.add_heading('6.2 Hitos y criterios de salida', level=2)
    doc.add_paragraph("• Hito 1 (Fin S3): MVP Operativo. Plataforma desplegada con carga controlada y módulos core habilitados.", style='List Bullet')
    doc.add_paragraph("• Hito 2 (Fin S4): Integraciones y Estabilización. Conectores habilitados y validación de datos.", style='List Bullet')
    doc.add_paragraph("• Hito 3 (Fin S5): Cierre. Aceptación de usuario (UAT) y entrega de documentación.", style='List Bullet')

    # 7. Entregables
    doc.add_heading('7. Entregables', level=1)
    
    doc.add_heading('7.1 Entregables Etapa 1 — MVP (S1–S3)', level=2)
    doc.add_paragraph("• Plataforma analítica MVP operativa.", style='List Bullet')
    doc.add_paragraph("• Modelo de datos y catálogo de KPIs.", style='List Bullet')
    doc.add_paragraph("• Módulo Inteligencia Comercial (Dashboards).", style='List Bullet')
    doc.add_paragraph("• Skava AI habilitado sobre dataset MVP.", style='List Bullet')
    
    doc.add_heading('7.2 Entregables Etapa 2 — Deploy y cierre (S4–S5)', level=2)
    doc.add_paragraph("• Integraciones ETL/ELT con sistemas productivos.", style='List Bullet')
    doc.add_paragraph("• Informe de validación y reconciliación de datos.", style='List Bullet')
    doc.add_paragraph("• Documentación técnica y Runbooks operacionales.", style='List Bullet')
    doc.add_paragraph("• Capacitación realizada.", style='List Bullet')
    
    doc.add_heading('7.3 Mantención mensual (post-entrega)', level=2)
    doc.add_paragraph("Servicio de soporte funcional y analítico en horario hábil (L-V), gestión de incidencias, monitoreo de operación y cobertura de costos de infraestructura bajo política de uso razonable.")

    # 8. Roles y responsabilidades
    doc.add_heading('8. Roles y responsabilidades', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Actividad"
    table.rows[0].cells[1].text = "Proveedor"
    table.rows[0].cells[2].text = "Cliente (Skava)"
    
    data = [
        ("Gestión de proyecto", "Responsable", "Product Owner / Sponsor"),
        ("Definición de KPIs", "Facilitador", "Dueños de KPI"),
        ("Desarrollo e Implementación", "Ejecutor", "Validación Técnica")
    ]
    for i, (act, prov, cli) in enumerate(data, 1):
        table.rows[i].cells[0].text = act
        table.rows[i].cells[1].text = prov
        table.rows[i].cells[2].text = cli

    # 9. Supuestos, dependencias y riesgos
    doc.add_heading('9. Supuestos, dependencias y riesgos', level=1)
    
    doc.add_heading('9.1 Supuestos y dependencias', level=2)
    doc.add_paragraph("• Disponibilidad de datos históricos (≥12 meses) para reconciliación en Etapa 2.", style='List Bullet')
    doc.add_paragraph("• Acceso a credenciales y documentación de APIs/BD para integraciones [PENDIENTE DE DEFINIR].", style='List Bullet')
    doc.add_paragraph("• Los costos se expresan en UF más IVA.", style='List Bullet')
    
    doc.add_heading('9.3 Riesgos técnicos y mitigaciones', level=2)
    doc.add_paragraph("• Calidad de datos origen: Se mitiga con reglas de DQ en Etapa 1 y acuerdos de tolerancia.", style='List Bullet')
    doc.add_paragraph("• Ambigüedad en KPIs: Se mitiga mediante catálogo versionado y firma de aceptación.", style='List Bullet')

    # 10. Soporte, SLA y continuidad
    doc.add_heading('10. Soporte, SLA y continuidad', level=1)
    doc.add_paragraph("• Horario base: Lunes a Viernes, horario comercial.", style='List Bullet')
    doc.add_paragraph("• Niveles de Severidad: P1 (Crítico) = 4h tiempo de respuesta; P2 (Alto) = 24h.", style='List Bullet')
    doc.add_paragraph("• No incluye soporte 24/7 salvo acuerdo específico.", style='List Bullet')

    # 11. Términos comerciales (alto nivel)
    doc.add_heading('11. Términos comerciales (alto nivel)', level=1)
    doc.add_paragraph("A continuación se detallan los valores de inversión para el proyecto:")
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    items = [
        ("Etapa 1 – MVP (S1–S3)", "UF 105 (Único / Setup)"),
        ("Etapa 2 – Deploy y cierre (S4–S5)", "UF 50 (Único / Setup)"),
        ("Total desarrollo", "UF 155 (+ IVA)"),
        ("Descuento cliente referente", "UF -155 (+ IVA)"),
        ("Total a pagar por desarrollo", "UF 0"),
        ("Mantención mensual (Soporte L-V)", "UF 15,5 (+ IVA)")
    ]
    
    for i, (desc, val) in enumerate(items):
        table.rows[i].cells[0].text = desc
        table.rows[i].cells[1].text = val
        
    doc.add_paragraph("\nNota: La referencia interna para planificación es 1 Story Point = 2 UF.")

    set_footer(doc)
    doc.save('Propuesta_Tecnica_Skava.docx')

if __name__ == "__main__":
    create_proposal_docx()